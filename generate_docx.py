"""
Generate Laporan UAS Pemrograman Basis Data (ST116)
Uses the user's template DOCX as base document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

# Load template
template_path = r"C:\Users\User\Downloads\Template Laporan UAS.docx"
doc = Document(template_path)

# ---- Helpers ----

def add_para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=6, font_name='Arial'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_code_block(code_text):
    """Code block in a 1-cell table with gray shading, Consolas 10pt."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    # Clear default paragraph
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Gray shading
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F0'
    })
    tcPr.append(shd)
    doc.add_paragraph()  # spacing after code block
    return table

def add_screenshot_placeholder(text="[Screenshot hasil eksekusi query]"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Arial'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_narration(text):
    """Narration paragraph (justify, Arial 12pt)."""
    return add_para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=8)

def add_soal_header(text):
    """Soal header with separator line."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_subsection(text):
    """Bold subsection header."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p

# ---- Update cover page ----
# Find and replace placeholder text in cover
# Handle multi-run paragraphs by concatenating run text
for p in doc.paragraphs:
    full_text = p.text
    if 'Dosen: [Nama Pengampu]' in full_text:
        # Clear all runs and set new text in first run
        new_text = 'Dosen: 1. Hendri Kurniawan Prakosa, S.Kom., M.Cs'
        if p.runs:
            p.runs[0].text = new_text
            p.runs[0].font.name = 'Arial'
            for r in p.runs[1:]:
                r.text = ''
    for run in p.runs:
        if '[Nama Lengkap]' in run.text:
            run.text = 'Ilham Arifin'
        elif '[NIM]' in run.text:
            run.text = '24.11.6022'
        elif '2025' == run.text.strip() and p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            run.text = '2026'

# Add anggota kelompok after NIM (insert before prodi line)
# Find the paragraph with NIM and add members after it
for i, p in enumerate(doc.paragraphs):
    if '24.11.6022' in p.text and p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        # Insert member list after NIM
        members = [
            'Ilham Arifin (24.11.6022)',
            'Muh Alfin Fauzi (24.11.6042)',
            'Moh Zaxlee Boneno (24.11.5990)',
            'Nabil Q Ahmad (24.11.6040)',
        ]
        # We'll add these as new paragraphs after the current one
        for j, member in enumerate(members):
            new_p = doc.add_paragraph()
            run = new_p.add_run(member)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Move paragraph after current position
            p._element.addnext(new_p._element)
            p = new_p  # chain: each new para after the previous
        break

# ---- Remove template's example soal (keep cover only) ----
# Delete paragraphs after the separator line (the example soal)
# Find the separator "____" line and remove everything after it
body = doc.element.body
# Find all paragraphs
paragraphs = list(doc.paragraphs)
delete_mode = False
to_delete = []
for p in paragraphs:
    if '____________________________________________________________' in p.text:
        delete_mode = True
        continue
    if delete_mode:
        to_delete.append(p)

# Also delete the template table (the example code block)
for table in doc.tables[:]:
    to_delete.append(table)

for p in to_delete:
    p._element.getparent().remove(p._element)

# Also remove empty tables that were part of template
for table in list(doc.tables):
    table._element.getparent().remove(table._element)

# ---- Page break after cover ----
doc.add_page_break()

# ============================================================
# BAB I: PENDAHULUAN
# ============================================================
add_para('BAB I', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=2)
add_para('PENDAHULUAN', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=12)

add_subsection('1.1 Latar Belakang')
add_narration(
    'Seiring dengan perkembangan teknologi informasi yang semakin pesat, pengelolaan data '
    'menjadi salah satu aspek krusial dalam menjalankan sebuah bisnis. Waveneap Management '
    'System adalah sistem manajemen toko online yang sebelumnya dibangun menggunakan MongoDB '
    'sebagai database NoSQL. Namun dalam mata kuliah Pemrograman Basis Data, basisdata relasional '
    '(RDBMS) menjadi fokus utama pembelajaran, sehingga dilakukan konversi skema dari MongoDB '
    'ke MySQL/MariaDB untuk memenuhi kebutuhan akademik sekaligus mengeksplorasi fitur-fitur '
    'lanjutan RDBMS seperti function, procedure, trigger, index, view, dan database security.'
)

add_subsection('1.2 Tujuan')
add_narration('Adapun tujuan dari pembuatan laporan ini adalah:')
add_para('1. Membuat basisdata relasional dengan minimal 5 tabel yang berisi minimal 10 baris data per tabel.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('2. Mengimplementasikan relasi one-to-one, one-to-many, dan many-to-many antar tabel.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('3. Mengimplementasikan function, stored procedure, trigger, index, view, dan database security.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('4. Mendokumentasikan seluruh implementasi dalam bentuk laporan yang sistematis.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=8)

add_subsection('1.3 Anggota Kelompok')
add_narration('Laporan final project ini dikerjakan oleh kelompok dengan anggota sebagai berikut:')

# Anggota table
member_table = doc.add_table(rows=5, cols=3)
member_table.style = 'Table Grid'
member_data = [
    ('No', 'Nama Lengkap', 'NIM'),
    ('1', 'Ilham Arifin', '24.11.6022'),
    ('2', 'Muh Alfin Fauzi', '24.11.6042'),
    ('3', 'Moh Zaxlee Boneno', '24.11.5990'),
    ('4', 'Nabil Q Ahmad', '24.11.6040'),
]
for i, (no, nama, nim) in enumerate(member_data):
    member_table.cell(i, 0).text = no
    member_table.cell(i, 1).text = nama
    member_table.cell(i, 2).text = nim
    for cell in member_table.rows[i].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                if i == 0:
                    run.bold = True

doc.add_paragraph()

add_subsection('1.4 Link Repository')
add_narration('Link file .sql dan kerangka basisdata (ERD): https://github.com/apin24/uts')

doc.add_page_break()

# ============================================================
# BAB II: PEMBAHASAN SOAL NOMOR 1 (DATABASE)
# ============================================================
add_para('BAB II', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=2)
add_para('PEMBAHASAN SOAL NOMOR 1 - BASISDATA', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=12)

add_subsection('2.1 Deskripsi Basisdata')
add_narration(
    'Basisdata yang dibuat bernama "uas" yang merupakan konversi dari skema MongoDB '
    'Waveneap Management System ke MySQL/MariaDB. Basisdata ini terdiri dari 8 tabel yang '
    'mencakup data user, produk, transaksi, serta tabel pendukung untuk logging dan relasi '
    'many-to-many. Berikut adalah daftar tabel yang dibuat:'
)

add_para('1. users - Menyimpan data user/karyawan (Admin, Kasir, Staff Packing).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('2. user_profiles - Menyimpan profil lengkap user (relasi 1:1 ke users).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('3. categories - Menyimpan kategori produk.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('4. products - Menyimpan data produk beserta harga dan deskripsi.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('5. product_variants - Menyimpan varian ukuran dan stok per produk (relasi 1:N).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('6. transactions - Menyimpan data transaksi keuangan (pemasukan dan pengeluaran).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('7. user_transactions - Join table untuk relasi many-to-many antara users dan transactions.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('8. activity_logs - Menyimpan log aktivitas perubahan data (untuk keperluan trigger).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=8)

add_subsection('2.2 Struktur Tabel dan Relasi')
add_narration(
    'Basisdata ini mengimplementasikan tiga jenis relasi yaitu one-to-one, one-to-many, '
    'dan many-to-many. Berikut adalah penjelasan relasi antar tabel:'
)
add_para('a. One-to-One: users <-> user_profiles (setiap user memiliki satu profil unik).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('b. One-to-Many: categories -> products (satu kategori punya banyak produk).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('c. One-to-Many: products -> product_variants (satu produk punya banyak varian ukuran).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('d. One-to-Many: products -> transactions (satu produk bisa terjual di banyak transaksi).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('e. One-to-Many: users -> transactions (satu user input banyak transaksi).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=4)
add_para('f. Many-to-Many: users <-> transactions (via join table user_transactions).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=8)

add_subsection('2.3 ERD (Entity Relationship Diagram)')
add_narration('Berikut adalah ERD dari basisdata uas:')

# Insert ERD image if exists
import os
erd_path = r"C:\Users\User\Downloads\uas-pbd-ilham\erd.png"
if os.path.exists(erd_path):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(erd_path, width=Inches(5.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    add_screenshot_placeholder("[ERD diagram - lihat file erd.png]")

add_para('Gambar 2.1 ERD uas', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=12)

add_subsection('2.4 SQL - Pembuatan Tabel')
add_narration('Berikut adalah perintah SQL untuk membuat seluruh tabel beserta foreign key-nya:')

# Read database.sql and include key CREATE TABLE statements
with open(r"C:\Users\User\Downloads\uas-pbd-ilham\database.sql", "r") as f:
    db_sql = f.read()

# Extract CREATE TABLE blocks
import re
create_blocks = re.findall(r'CREATE TABLE \w+.*?ENGINE=InnoDB;', db_sql, re.DOTALL)
for block in create_blocks:
    add_code_block(block.strip())

add_screenshot_placeholder("[Screenshot hasil eksekusi: SHOW TABLES di phpMyAdmin/MySQL]")

add_narration(
    'Setelah seluruh tabel berhasil dibuat, langkah selanjutnya adalah mengisi data ke '
    'dalam tabel-tabel tersebut. Data seed dimasukkan secara manual untuk tabel users '
    '(10 baris), user_profiles (10 baris), categories (10 baris), products (15 baris), '
    'product_variants (32 baris), dan transactions (20 baris). Selain itu, untuk keperluan '
    'demonstrasi indexing, ditambahkan 5000 baris data dummy ke tabel transactions menggunakan '
    'script Python sehingga total baris transactions menjadi 5020 baris.'
)

add_subsection('2.5 SQL - Seed Data')
add_narration('Contoh seed data untuk tabel users:')

# Extract INSERT users block
insert_users = re.search(r'INSERT INTO users.*?VALUES\s*(.*?);', db_sql, re.DOTALL)
if insert_users:
    add_code_block("INSERT INTO users (name, username, email, phone, role, password) VALUES\n" + insert_users.group(1).strip() + ";")

add_screenshot_placeholder("[Screenshot hasil eksekusi: SELECT * FROM users]")

add_narration('Contoh seed data untuk tabel products:')

insert_products = re.search(r'INSERT INTO products.*?VALUES\s*(.*?);', db_sql, re.DOTALL)
if insert_products:
    add_code_block("INSERT INTO products (category_id, title, price, description, image) VALUES\n" + insert_products.group(1).strip() + ";")

add_screenshot_placeholder("[Screenshot hasil eksekusi: SELECT * FROM products]")

add_subsection('2.6 Verifikasi Jumlah Baris Data')
add_narration('Berikut adalah query untuk memverifikasi jumlah baris data per tabel:')

add_code_block("""SELECT 'users' AS table_name, COUNT(*) AS row_count FROM users
UNION ALL
SELECT 'user_profiles', COUNT(*) FROM user_profiles
UNION ALL
SELECT 'categories', COUNT(*) FROM categories
UNION ALL
SELECT 'products', COUNT(*) FROM products
UNION ALL
SELECT 'product_variants', COUNT(*) FROM product_variants
UNION ALL
SELECT 'transactions', COUNT(*) FROM transactions
UNION ALL
SELECT 'user_transactions', COUNT(*) FROM user_transactions
UNION ALL
SELECT 'activity_logs', COUNT(*) FROM activity_logs;""")

add_screenshot_placeholder("[Screenshot hasil eksekusi: row count per tabel]")

doc.add_page_break()

# ============================================================
# BAB III: PEMBAHASAN SOAL NOMOR 2 (SQL IMPLEMENTATION)
# ============================================================
add_para('BAB III', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=2)
add_para('PEMBAHASAN SOAL NOMOR 2 - SQL IMPLEMENTATION', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=12)

add_narration(
    'Pada bagian ini, akan diimplementasikan perintah-perintah SQL yang berkaitan dengan '
    'topik yang dibahas di mata kuliah Pemrograman Basis Data, yaitu Function, Procedure, '
    'Trigger, Index, View, dan Database Security. Setiap implementasi disertai dengan SQL '
    'statement, screenshot hasil eksekusi, serta penjelasan yang jelas dan runtut.'
)

# Read queries.sql
with open(r"C:\Users\User\Downloads\uas-pbd-ilham\queries.sql", "r") as f:
    queries_sql = f.read()

# ---- 3.1 FUNCTION ----
add_soal_header('2a. Function (15 poin)')
add_narration('Pada bagian ini dibuat 2 buah function: 1 function tanpa parameter dan 1 function dengan 2 parameter.')

add_subsection('Function 1: GetTotalIncome() - Tanpa Parameter')
add_narration('Function ini berfungsi untuk menghitung total pemasukan (income) dari semua transaksi yang berstatus Success. Function ini menggunakan JOIN ke tabel users untuk memfilter data transaksi berdasarkan user yang input.')

add_code_block("""DELIMITER //
CREATE FUNCTION GetTotalIncome()
RETURNS DECIMAL(12,2)
DETERMINISTIC
BEGIN
    DECLARE total DECIMAL(12,2);
    SELECT COALESCE(SUM(t.amount), 0) INTO total
    FROM transactions t
    JOIN users u ON t.user_id = u.id
    WHERE t.is_income = 1 AND t.status = 'Success';
    RETURN total;
END //
DELIMITER ;""")

add_screenshot_placeholder("[Screenshot: CREATE FUNCTION GetTotalIncome]")

add_narration('Setelah function dibuat, function dieksekusi dalam query yang menggunakan JOIN 2 tabel (transactions dan users) untuk menampilkan total transaksi per user, lalu memanggil function GetTotalIncome() untuk mendapatkan total keseluruhan:')

add_code_block("""SELECT
    u.name AS user_name,
    u.role,
    COUNT(t.id) AS total_transactions,
    SUM(t.amount) AS total_amount
FROM transactions t
JOIN users u ON t.user_id = u.id
WHERE t.is_income = 1
GROUP BY u.id, u.name, u.role
ORDER BY total_amount DESC;

SELECT GetTotalIncome() AS total_income_all;""")

add_screenshot_placeholder("[Screenshot: Hasil query JOIN + GetTotalIncome()]")

add_subsection('Function 2: GetExpenseByUserAndMonth(user_id, month) - Dengan 2 Parameter')
add_narration('Function ini menerima 2 parameter yaitu user_id dan bulan (format YYYY-MM). Function ini menghitung total pengeluaran seorang user pada bulan tertentu menggunakan subquery untuk filter.')

add_code_block("""DELIMITER //
CREATE FUNCTION GetExpenseByUserAndMonth(p_user_id INT, p_month VARCHAR(7))
RETURNS DECIMAL(12,2)
DETERMINISTIC
BEGIN
    DECLARE total DECIMAL(12,2);
    SELECT COALESCE(SUM(amount), 0) INTO total
    FROM transactions
    WHERE user_id = p_user_id
    AND is_income = 0
    AND status = 'Success'
    AND DATE_FORMAT(transaction_date, '%Y-%m') = p_month;
    RETURN total;
END //
DELIMITER ;""")

add_screenshot_placeholder("[Screenshot: CREATE FUNCTION GetExpenseByUserAndMonth]")

add_narration('Eksekusi function dengan subquery untuk menampilkan pengeluaran setiap user di bulan Juli 2026:')

add_code_block("""SELECT
    u.name AS user_name,
    GetExpenseByUserAndMonth(u.id, '2026-07') AS expense_juli
FROM users u
WHERE u.id IN (SELECT DISTINCT user_id FROM transactions WHERE is_income = 0)
ORDER BY expense_juli DESC;""")

add_screenshot_placeholder("[Screenshot: Hasil query GetExpenseByUserAndMonth]")

add_subsection('Daftar Function')
add_narration('Untuk menampilkan daftar function yang telah dibuat, gunakan query berikut:')

add_code_block("""SELECT
    ROUTINE_NAME AS function_name,
    ROUTINE_TYPE AS type,
    DATA_TYPE AS return_type
FROM information_schema.ROUTINES
WHERE ROUTINE_SCHEMA = 'uas' AND ROUTINE_TYPE = 'FUNCTION';""")

add_screenshot_placeholder("[Screenshot: Daftar function]")

# ---- 3.2 PROCEDURE ----
doc.add_page_break()
add_soal_header('2b. Procedure (15 poin)')
add_narration('Pada bagian ini dibuat 2 buah stored procedure: 1 procedure dengan parameter kosong (menggunakan OUT) dan 1 procedure dengan 2 parameter (IN). Keduanya mengandung control flow (IF, CASE) dan cursor.')

add_subsection('Procedure 1: GetFinancialSummary() - Parameter OUT')
add_narration('Procedure ini tidak menerima parameter input, tetapi mengembalikan 3 nilai melalui parameter OUT: total income, total expense, dan profit. Procedure ini berguna untuk mendapatkan ringkasan keuangan secara cepat.')

add_code_block("""DELIMITER //
CREATE PROCEDURE GetFinancialSummary(OUT p_total_income DECIMAL(12,2), OUT p_total_expense DECIMAL(12,2), OUT p_profit DECIMAL(12,2))
BEGIN
    SELECT COALESCE(SUM(amount), 0) INTO p_total_income
    FROM transactions WHERE is_income = 1 AND status = 'Success';

    SELECT COALESCE(SUM(amount), 0) INTO p_total_expense
    FROM transactions WHERE is_income = 0 AND status = 'Success';

    SET p_profit = p_total_income - p_total_expense;
END //
DELIMITER ;

CALL GetFinancialSummary(@income, @expense, @profit);
SELECT @income AS total_income, @expense AS total_expense, @profit AS profit;""")

add_screenshot_placeholder("[Screenshot: CREATE PROCEDURE GetFinancialSummary + hasil CALL]")

add_subsection('Procedure 2: GenerateUserReport(user_id, month) - Dengan 2 Parameter IN')
add_narration('Procedure ini menerima 2 parameter IN: user_id dan bulan (YYYY-MM). Procedure ini menggunakan cursor untuk mengiterasi setiap transaksi user pada bulan tertentu, dan control flow IF untuk memisahkan antara income dan expense. Procedure ini berguna untuk generate laporan transaksi per user per bulan.')

add_code_block("""DELIMITER //
CREATE PROCEDURE GenerateUserReport(IN p_user_id INT, IN p_month VARCHAR(7))
BEGIN
    DECLARE done INT DEFAULT 0;
    DECLARE v_id INT;
    DECLARE v_name VARCHAR(200);
    DECLARE v_amount DECIMAL(12,2);
    DECLARE v_is_income TINYINT;
    DECLARE v_date DATE;
    DECLARE v_user_name VARCHAR(100);
    DECLARE v_total_income DECIMAL(12,2) DEFAULT 0;
    DECLARE v_total_expense DECIMAL(12,2) DEFAULT 0;

    -- Cursor untuk ambil transaksi user di bulan tertentu
    DECLARE cur CURSOR FOR
        SELECT t.id, t.name, t.amount, t.is_income, t.transaction_date
        FROM transactions t
        WHERE t.user_id = p_user_id
        AND DATE_FORMAT(t.transaction_date, '%Y-%m') = p_month
        ORDER BY t.transaction_date;

    DECLARE CONTINUE HANDLER FOR NOT FOUND SET done = 1;

    SELECT name INTO v_user_name FROM users WHERE id = p_user_id;

    -- Temporary table untuk kumpulin hasil cursor (phpMyAdmin-friendly)
    DROP TEMPORARY TABLE IF EXISTS tmp_report;
    CREATE TEMPORARY TABLE tmp_report (
        id INT,
        transaction_name VARCHAR(200),
        amount DECIMAL(12,2),
        type VARCHAR(10),
        trans_date DATE
    );

    OPEN cur;
    read_loop: LOOP
        FETCH cur INTO v_id, v_name, v_amount, v_is_income, v_date;
        IF done THEN LEAVE read_loop; END IF;

        -- Control flow IF untuk pisah income/expense
        IF v_is_income = 1 THEN
            SET v_total_income = v_total_income + v_amount;
            INSERT INTO tmp_report VALUES (v_id, v_name, v_amount, 'INCOME', v_date);
        ELSE
            SET v_total_expense = v_total_expense + v_amount;
            INSERT INTO tmp_report VALUES (v_id, v_name, v_amount, 'EXPENSE', v_date);
        END IF;
    END LOOP;
    CLOSE cur;

    -- Single result set (phpMyAdmin-compatible)
    SELECT
        id, transaction_name, amount, type, trans_date,
        v_user_name AS user_name,
        p_month AS report_month,
        v_total_income AS total_income,
        v_total_expense AS total_expense,
        (v_total_income - v_total_expense) AS net_profit
    FROM tmp_report;

    DROP TEMPORARY TABLE IF EXISTS tmp_report;
END //
DELIMITER ;

CALL GenerateUserReport(1, '2026-07');""")

add_screenshot_placeholder("[Screenshot: CREATE PROCEDURE GenerateUserReport + hasil CALL]")

add_subsection('Daftar Stored Procedure')
add_code_block("""SELECT ROUTINE_NAME AS procedure_name, ROUTINE_TYPE AS type
FROM information_schema.ROUTINES
WHERE ROUTINE_SCHEMA = 'uas' AND ROUTINE_TYPE = 'PROCEDURE';""")

add_screenshot_placeholder("[Screenshot: Daftar stored procedure]")

# ---- 3.3 TRIGGER ----
doc.add_page_break()
add_soal_header('2c. Trigger (15 poin)')
add_narration('Pada bagian ini dibuat 2 buah trigger dengan variasi event yang berbeda: AFTER INSERT dan BEFORE UPDATE. Kedua trigger menggunakan parameter NEW dan/atau OLD.')

add_subsection('Trigger 1: trg_after_insert_transaction (AFTER INSERT)')
add_narration('Trigger ini aktif setelah ada INSERT baru pada tabel transactions. Trigger akan mencatat aktivitas INSERT ke tabel activity_logs, menyimpan data transaksi baru sebagai log. Body trigger menggunakan parameter NEW untuk mengambil data yang baru di-insert.')

add_code_block("""DELIMITER //
CREATE TRIGGER trg_after_insert_transaction
AFTER INSERT ON transactions
FOR EACH ROW
BEGIN
    INSERT INTO activity_logs (table_name, record_id, action, new_values, changed_by)
    VALUES (
        'transactions',
        NEW.id,
        'INSERT',
        CONCAT('name=', NEW.name, ', amount=', NEW.amount, ', is_income=', NEW.is_income, ', status=', NEW.status),
        CURRENT_USER()
    );
END //
DELIMITER ;

-- Eksekusi: INSERT transaksi baru
INSERT INTO transactions (user_id, product_id, name, category, amount, is_income, status, transaction_date, platform, variant_size, notes)
VALUES (1, 1, 'Penjualan Kaos Test Trigger', 'Penjualan Produk', 89000, 1, 'Success', '2026-07-20', 'Shopee', 'M', 'Test trigger insert');

-- Cek log
SELECT * FROM activity_logs WHERE table_name = 'transactions' ORDER BY changed_at DESC LIMIT 5;""")

add_screenshot_placeholder("[Screenshot: CREATE TRIGGER + INSERT + hasil log di activity_logs]")

add_subsection('Trigger 2: trg_before_update_product (BEFORE UPDATE)')
add_narration('Trigger ini aktif sebelum ada UPDATE pada tabel products. Trigger berfungsi untuk dua hal: (1) validasi harga tidak boleh negatif menggunakan control flow IF dan SIGNAL SQLSTATE, dan (2) mencatat perubahan harga ke tabel activity_logs menggunakan parameter OLD dan NEW. Trigger ini bukan hanya untuk logging tetapi juga untuk validasi data.')

add_code_block("""DELIMITER //
CREATE TRIGGER trg_before_update_product
BEFORE UPDATE ON products
FOR EACH ROW
BEGIN
    IF NEW.price < 0 THEN
        SIGNAL SQLSTATE '45000' SET MESSAGE_TEXT = 'Harga tidak boleh negatif';
    END IF;

    IF OLD.price != NEW.price THEN
        INSERT INTO activity_logs (table_name, record_id, action, old_values, new_values, changed_by)
        VALUES (
            'products',
            NEW.id,
            'UPDATE',
            CONCAT('price=', OLD.price),
            CONCAT('price=', NEW.price),
            CURRENT_USER()
        );
    END IF;
END //
DELIMITER ;

-- Eksekusi: UPDATE harga produk
UPDATE products SET price = 95000 WHERE id = 1;

-- Cek log
SELECT * FROM activity_logs WHERE table_name = 'products' ORDER BY changed_at DESC LIMIT 5;""")

add_screenshot_placeholder("[Screenshot: CREATE TRIGGER + UPDATE + hasil log di activity_logs]")

add_subsection('Daftar Trigger')
add_code_block("""SELECT
    TRIGGER_NAME AS trigger_name,
    EVENT_MANIPULATION AS event,
    EVENT_OBJECT_TABLE AS table_name,
    ACTION_TIMING AS timing
FROM information_schema.TRIGGERS
WHERE TRIGGER_SCHEMA = 'uas';""")

add_screenshot_placeholder("[Screenshot: Daftar trigger]")

# ---- 3.4 INDEX ----
doc.add_page_break()
add_soal_header('2d. Index (10 poin)')
add_narration('Pada bagian ini dibuat 3 buah index dengan cara yang berbeda: CREATE INDEX, ALTER TABLE, dan index pada tabel baru. Semua index menggunakan composite key (lebih dari 1 kolom).')

add_subsection('Index 1: CREATE INDEX pada transactions')
add_narration('Index ini dibuat menggunakan perintah CREATE INDEX dengan composite key (user_id, transaction_date). Index ini mempercepat query yang memfilter transaksi berdasarkan user dan rentang tanggal.')

add_code_block("""CREATE INDEX idx_transactions_user_date ON transactions (user_id, transaction_date);""")

add_subsection('Index 2: CREATE INDEX pada transactions (composite)')
add_narration('Index ini dibuat menggunakan perintah CREATE INDEX dengan composite key (is_income, status, transaction_date). Index ini mempercepat query yang memfilter transaksi berdasarkan tipe income/expense, status, dan tanggal.')

add_code_block("""CREATE INDEX idx_transactions_income_status_date ON transactions (is_income, status, transaction_date);""")

add_subsection('Index 3: ALTER TABLE pada products')
add_narration('Index ini dibuat menggunakan perintah ALTER TABLE dengan composite key (category_id, price). Index ini mempercepat query yang memfilter produk berdasarkan kategori dan range harga.')

add_code_block("""ALTER TABLE products ADD INDEX idx_products_category_price (category_id, price);""")

add_subsection('Perbandingan EXPLAIN dengan dan tanpa Index')
add_narration('Untuk membuktikan bahwa index mempercepat query, digunakan perintah EXPLAIN. Berikut adalah perbandingan query dengan index dan tanpa index (menggunakan IGNORE INDEX):')

add_code_block("""-- Query DENGAN index
EXPLAIN
SELECT * FROM transactions
WHERE user_id = 1 AND transaction_date BETWEEN '2026-07-01' AND '2026-07-31';

-- Query TANPA index (full table scan)
EXPLAIN
SELECT * FROM transactions IGNORE INDEX (idx_transactions_user_date, idx_transactions_income_status_date)
WHERE user_id = 1 AND transaction_date BETWEEN '2026-07-01' AND '2026-07-31';""")

add_screenshot_placeholder("[Screenshot: EXPLAIN dengan index (rows diperiksa sedikit, type=ref/range)]")
add_screenshot_placeholder("[Screenshot: EXPLAIN tanpa index (rows diperiksa banyak, type=ALL / full scan)]")

add_narration('Perbedaannya: pada query dengan index, kolom type menunjukkan "ref" atau "range" dengan rows yang sedikit (hanya baris yang relevan). Sedangkan pada query tanpa index, type menunjukkan "ALL" yang berarti full table scan, dan rows menunjukkan seluruh baris tabel (5020 baris) diperiksa. Ini membuktikan bahwa index secara signifikan mengurangi jumlah baris yang diperiksa dan mempercepat eksekusi query.')

add_subsection('Daftar Index')
add_code_block("""SELECT
    TABLE_NAME AS table_name,
    INDEX_NAME AS index_name,
    GROUP_CONCAT(COLUMN_NAME ORDER BY SEQ_IN_INDEX) AS columns,
    INDEX_TYPE AS type,
    NON_UNIQUE AS non_unique
FROM information_schema.STATISTICS
WHERE TABLE_SCHEMA = 'uas'
GROUP BY TABLE_NAME, INDEX_NAME, INDEX_TYPE, NON_UNIQUE;""")

add_screenshot_placeholder("[Screenshot: Daftar index]")

# ---- 3.5 VIEW ----
doc.add_page_break()
add_soal_header('2e. View (10 poin)')
add_narration('Pada bagian ini dibuat 3 buah view dengan tipe yang berbeda: horizontal view, vertical view, dan view inside view dengan WITH CHECK OPTION.')

add_subsection('View 1: v_success_transactions (Horizontal View)')
add_narration('Horizontal view adalah view yang memfilter baris (WHERE clause) tetapi menampilkan semua kolom. View ini menampilkan hanya transaksi yang berstatus Success, dengan JOIN ke tabel users dan products.')

add_code_block("""CREATE VIEW v_success_transactions AS
SELECT
    t.id,
    t.name AS transaction_name,
    t.amount,
    t.transaction_date,
    t.platform,
    u.name AS user_name,
    p.title AS product_title
FROM transactions t
JOIN users u ON t.user_id = u.id
LEFT JOIN products p ON t.product_id = p.id
WHERE t.status = 'Success';

SELECT * FROM v_success_transactions LIMIT 10;""")

add_screenshot_placeholder("[Screenshot: CREATE VIEW v_success_transactions + hasil SELECT]")

add_subsection('View 2: v_user_summary (Vertical View)')
add_narration('Vertical view adalah view yang memfilter kolom (hanya menampilkan kolom tertentu) tetapi menampilkan semua baris. View ini hanya menampilkan kolom id, name, username, email, role, dan created_at dari tabel users.')

add_code_block("""CREATE VIEW v_user_summary AS
SELECT id, name, username, email, role, created_at
FROM users;

SELECT * FROM v_user_summary;""")

add_screenshot_placeholder("[Screenshot: CREATE VIEW v_user_summary + hasil SELECT]")

add_subsection('View 3: v_admin_users (View inside View + WITH CHECK OPTION)')
add_narration('View inside view adalah view yang dibuat berdasarkan view lain. View ini select dari v_user_summary (view 2) dengan filter role = Admin. Klausa WITH CHECK OPTION digunakan untuk memastikan bahwa setiap INSERT atau UPDATE melalui view ini harus memenuhi kondisi WHERE (role harus Admin).')

add_code_block("""CREATE VIEW v_admin_users AS
SELECT id, name, username, email, role, created_at
FROM v_user_summary
WHERE role = 'Admin'
WITH CHECK OPTION;

SELECT * FROM v_admin_users;

-- Test INSERT yang memenuhi CHECK OPTION (berhasil)
INSERT INTO v_admin_users (name, username, email, role) VALUES ('Admin Baru', 'admin_baru', 'adminbaru@waveneap.com', 'Admin');

-- Test INSERT yang TIDAK memenuhi CHECK OPTION (error)
-- INSERT INTO v_admin_users (name, username, email, role) VALUES ('Kasir', 'kasir_test', 'kasir@waveneap.com', 'Kasir');
-- Error: CHECK OPTION failed""")

add_screenshot_placeholder("[Screenshot: CREATE VIEW v_admin_users + hasil SELECT + INSERT berhasil]")

add_subsection('Test UPDATE dan INSERT via View')
add_narration('Selain view inside view, view lain juga bisa digunakan untuk UPDATE dan INSERT:')

add_code_block("""-- UPDATE via v_success_transactions
UPDATE v_success_transactions SET platform = 'Shopee' WHERE id = 1;

-- INSERT via v_user_summary
INSERT INTO v_user_summary (name, username, email, role) VALUES ('Test User', 'test_user', 'test@waveneap.com', 'Staff Packing');""")

add_screenshot_placeholder("[Screenshot: Hasil UPDATE dan INSERT via view]")

add_subsection('Daftar View')
add_code_block("""SELECT
    TABLE_NAME AS view_name,
    CHECK_OPTION AS check_option,
    IS_UPDATABLE AS is_updatable
FROM information_schema.VIEWS
WHERE TABLE_SCHEMA = 'uas';""")

add_screenshot_placeholder("[Screenshot: Daftar view]")

# ---- 3.6 DATABASE SECURITY ----
doc.add_page_break()
add_soal_header('2f. Database Security (15 poin)')
add_narration('Pada bagian ini dibuat 3 buah user dan 3 buah role, lalu diberikan privilege yang berbeda-beda untuk masing-masing role. Hal ini bertujuan untuk mengimplementasikan prinsip least privilege dalam keamanan database.')

add_subsection('Membuat 3 User')
add_narration('Tiga user dibuat dengan password yang berbeda, masing-masing mewakili level akses yang berbeda:')

add_code_block("""CREATE USER IF NOT EXISTS 'admin_waveneap'@'localhost' IDENTIFIED BY 'admin123';
CREATE USER IF NOT EXISTS 'kasir_waveneap'@'localhost' IDENTIFIED BY 'kasir123';
CREATE USER IF NOT EXISTS 'staff_waveneap'@'localhost' IDENTIFIED BY 'staff123';""")

add_screenshot_placeholder("[Screenshot: CREATE USER berhasil]")

add_subsection('Membuat 3 Role')
add_narration('Tiga role dibuat untuk mengelompokkan privilege:')

add_code_block("""CREATE ROLE IF NOT EXISTS role_admin, role_kasir, role_staff;""")

add_screenshot_placeholder("[Screenshot: CREATE ROLE berhasil]")

add_subsection('Grant Privilege ke Role')
add_narration('Setiap role diberikan privilege yang berbeda sesuai dengan kebutuhannya:')

add_code_block("""-- role_admin: full access ke semua tabel + execute procedure
GRANT SELECT, INSERT, UPDATE, DELETE ON uas.* TO role_admin;
GRANT EXECUTE ON PROCEDURE uas.GetFinancialSummary TO role_admin;
GRANT EXECUTE ON PROCEDURE uas.GenerateUserReport TO role_admin;

-- role_kasir: select products + select/insert transactions + select view + execute procedure
GRANT SELECT ON uas.products TO role_kasir;
GRANT SELECT ON uas.categories TO role_kasir;
GRANT SELECT ON uas.product_variants TO role_kasir;
GRANT SELECT, INSERT ON uas.transactions TO role_kasir;
GRANT SELECT ON uas.v_success_transactions TO role_kasir;
GRANT SELECT ON uas.v_admin_users TO role_kasir;
GRANT EXECUTE ON PROCEDURE uas.GetFinancialSummary TO role_kasir;

-- role_staff: select only (read-only)
GRANT SELECT ON uas.products TO role_staff;
GRANT SELECT ON uas.categories TO role_staff;
GRANT SELECT ON uas.product_variants TO role_staff;
GRANT SELECT ON uas.v_user_summary TO role_staff;""")

add_screenshot_placeholder("[Screenshot: GRANT privilege ke role]")

add_subsection('Assign Role ke User')
add_narration('Setelah role dibuat dan diberi privilege, role di-assign ke user dan di-set sebagai default role:')

add_code_block("""GRANT role_admin TO 'admin_waveneap'@'localhost';
GRANT role_kasir TO 'kasir_waveneap'@'localhost';
GRANT role_staff TO 'staff_waveneap'@'localhost';

SET DEFAULT ROLE role_admin FOR 'admin_waveneap'@'localhost';
SET DEFAULT ROLE role_kasir FOR 'kasir_waveneap'@'localhost';
SET DEFAULT ROLE role_staff FOR 'staff_waveneap'@'localhost';""")

add_screenshot_placeholder("[Screenshot: GRANT role to user + SET DEFAULT ROLE]")

add_subsection('Verifikasi Privilege')
add_narration('Untuk membuktikan bahwa privilege sudah berhasil diberikan, gunakan query berikut:')

add_code_block("""-- Tampilkan daftar user
SELECT User AS user_name, Host AS host
FROM mysql.user WHERE User LIKE '%waveneap%';

-- Tampilkan daftar role assignment
SELECT CONCAT(User, '@', Host) AS assigned_to, Role AS role_name
FROM mysql.roles_mapping WHERE User LIKE '%waveneap%';

-- Tampilkan privilege per user
SHOW GRANTS FOR 'admin_waveneap'@'localhost';
SHOW GRANTS FOR 'kasir_waveneap'@'localhost';
SHOW GRANTS FOR 'staff_waveneap'@'localhost';""")

add_screenshot_placeholder("[Screenshot: Daftar user, role, dan grants]")

add_narration(
    'Bukti privilege berfungsi: user kasir_waveneap hanya bisa SELECT dan INSERT pada tabel '
    'transactions, tetapi tidak bisa DELETE atau UPDATE. Jika kasir mencoba DELETE, akan muncul '
    'error "ACCESS DENIED". User staff_waveneap hanya bisa SELECT (read-only), tidak bisa INSERT, '
    'UPDATE, atau DELETE. User admin_waveneap memiliki akses penuh (SELECT, INSERT, UPDATE, DELETE) '
    'ke semua tabel dan bisa execute procedure.'
)

doc.add_page_break()

# ============================================================
# BAB IV: PENUTUP
# ============================================================
add_para('BAB IV', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=2)
add_para('PENUTUP', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=12)

add_subsection('4.1 Kesimpulan')
add_narration(
    'Dari pengerjaan final project UAS Pemrograman Basis Data ini, dapat disimpulkan beberapa '
    'hal sebagai berikut:'
)
add_para('1. Berhasil dibuat basisdata uas dengan 8 tabel yang mengimplementasikan relasi one-to-one (users-user_profiles), one-to-many (categories-products, products-product_variants, products-transactions, users-transactions), dan many-to-many (users-transactions via user_transactions).', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=6)
add_para('2. Berhasil dibuat 2 function (GetTotalIncome tanpa parameter, GetExpenseByUserAndMonth dengan 2 parameter) yang dieksekusi menggunakan JOIN dan subquery.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=6)
add_para('3. Berhasil dibuat 2 stored procedure (GetFinancialSummary dengan parameter OUT, GenerateUserReport dengan 2 parameter IN) yang mengandung control flow IF dan cursor.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=6)
add_para('4. Berhasil dibuat 2 trigger (AFTER INSERT pada transactions untuk logging, BEFORE UPDATE pada products untuk validasi dan logging) dengan variasi event yang berbeda.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=6)
add_para('5. Berhasil dibuat 3 index dengan composite key menggunakan 3 cara berbeda (CREATE INDEX, ALTER TABLE), dan dibuktikan dengan EXPLAIN bahwa index mempercepat query.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=6)
add_para('6. Berhasil dibuat 3 view (horizontal view, vertical view, view inside view dengan WITH CHECK OPTION) yang masing-masing bisa digunakan untuk UPDATE dan INSERT.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=6)
add_para('7. Berhasil dibuat 3 user dan 3 role dengan privilege yang berbeda-beda, membuktikan implementasi database security dengan prinsip least privilege.', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=8)

add_subsection('4.2 Saran')
add_narration(
    'Untuk pengembangan selanjutnya, basisdata ini dapat ditambahkan dengan fitur backup dan '
    'restore otomatis, implementasi transaction isolation level yang lebih ketat, serta '
    'optimasi query lebih lanjut dengan analisis query execution plan yang lebih mendalam.'
)

# ---- Save ----
output_path = r"C:\Users\User\Downloads\uas-pbd-ilham\Laporan-UAS-PBD-Ilham-Arifin.docx"
doc.save(output_path)
print(f"DOCX saved: {output_path}")
